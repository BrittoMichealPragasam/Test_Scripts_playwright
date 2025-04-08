#*****************************************IM Master Data Table Onboarding Test Scenarios Automation******************************************************************************************
"""
    Business Objective: To validate the new SAP table onboarding process in GCP
    Description:
            The test cases will validate the new SAP table available in WDP as per the mappings provided and validate the data in the table is loaded into the GCP BigQuery tables
    Test Scope
            1. Verify Whether Tables and View Available in GCP as per mapping sheet
            2. Fetch Inbound Files from Bucket and validate dag run
            3. Verify Audit Table Validations
            4. Verify Data Loaded into Currated, History, View Tables
            5. Validate Primary Key/delta columns duplicates and Nulls
    Out of Scope:
            1. Fetching the count of SAP HANA and compare with GCP
            2. Validations in SAP system
            3. Config Validations - Insert, Update, Delete
    Authors:
            1. Britto Micheal Pragasam(bpragasam@tcs.woolworths.com.au)
        """
#****************************************************************************************************************************************************************************

#Import Required Libraries and Packages
import csv
import datetime
import pandas as pd
from playwright.sync_api import sync_playwright
import pytest
import os
from docx import Document
from docx.shared import Inches, RGBColor
from GCPLibrary import *
from GCP_BigQuery_Query_Library import *
from UC4_Informatica_Library import *

# Constants
SCREENSHOT_DIR = 'S:\\IM_Automation_Repositories\\05.TestResults\\02.Screenshots'
#Fetch inputs from Environment Variables
CSV_FILE_PATH = "S:\\IM_Automation_Repositories\\03.InputFileTemplates\\01.MasterData\\NewTableOnboarding.csv"
TEST_RESULT_PATH = "S:\\IM_Automation_Repositories\\03.InputFileTemplates\\01.MasterData"

#Get User ID
USER_ID = getUserId()

#Function to read CSV file and store the data in a list
def read_csv(file_path):
    with open(file_path, mode='r', newline='') as file:
        return list(csv.DictReader(file))

#Fixture to launch the browser
@pytest.fixture(scope="class")
def google_cloud_helper():
    with sync_playwright() as playwright:
        user_data_dir = f"C://Users//{USER_ID}//AppData//Local//Google//Chrome//User Data"
        helper = GoogleCloudHelper(playwright, user_data_dir)
        browser = helper.launch_browser()
        page = browser.new_page()
        browser.pages[0].close()
        yield helper, browser, page
        browser.close()

#Fixture to read the test data from the CSV file - It helps to run the test cases in each row of the CSV file
@pytest.fixture(scope="class", params=read_csv(CSV_FILE_PATH))
def test_data(request):
    return request.param
#Fixture to create a document to store the test results
@pytest.fixture(scope="class")
def document(test_data):
    doc = Document()
    yield doc
    doc_path = os.path.join(TEST_RESULT_PATH, f"test_report_{test_data['InboundFileName'].replace('.', '_')}.docx")
    os.makedirs(os.path.dirname(doc_path), exist_ok=True)
    doc.save(doc_path)

#Test Class to run the test cases
class TestGoogleCloudWorkflow:
    strJobStartTime = None
    strJobEndTime = None
    strNumberOfRecords = None
    strInboundFileName = None
    strUC4JobEndTime = None
    strActualProcessedFileName = None
    df_Mapped_FileData  = pd.DataFrame()
    def take_screenshot(self, page, test_data, test_name, document):
        screenshot_path = os.path.join(SCREENSHOT_DIR,test_data['InboundFileName'], f"{test_name}.png")
        os.makedirs(os.path.dirname(screenshot_path), exist_ok=True)
        page.screenshot(path=screenshot_path)
        document.add_heading(test_name, level=2)
        document.add_picture(screenshot_path, width=Inches(6))
    def test_login_Navigate_to_UC4(self, google_cloud_helper, test_data, document):
        helper, browser, page = google_cloud_helper
        if "." not in test_data['InboundFileName']:
            page.goto("https://awa-nonprod.woolworths.com.au/awi/UC4NONPD")
            if "uat" in test_data['ProjectId']:
                intClientID = "300"
            elif "test" in test_data['ProjectId']:
                intClientID = "200"
            UC4InformaticaHelper(helper.playwright).login_UC4_and_navigate_to_homepage(page, intClientID, "XHRN1", "Ashmi2021")
            self.take_screenshot(page, test_data, f'launch UC4', document)
    def test_search_UC4_Job_and_execute(self, google_cloud_helper, test_data,document):
        helper, browser, page = google_cloud_helper
        if "." not in test_data['InboundFileName']:
            strUC4JobStatus = UC4InformaticaHelper(helper.playwright).search_uc4_Job_and_execute(page, test_data['UC4JobName'])
            if strUC4JobStatus != "FAIL" and strUC4JobStatus is not None:
                self.take_screenshot(page, test_data,f'Search UC4 Job and Execute', document)
            else:
                self.take_screenshot(page, test_data, f'FAIL - Search UC4 Job and Execute', document)
    def test_verify_UC4_Job_execution_status(self, google_cloud_helper, test_data,document):
        helper, browser, page = google_cloud_helper
        if "." not in test_data['InboundFileName']:
            strUC4JobVerificationStatus, TestGoogleCloudWorkflow.strUC4JobEndTime = UC4InformaticaHelper(helper.playwright).verify_Uc4_Job_execution_status(page)
            if strUC4JobVerificationStatus != "FAIL" and strUC4JobVerificationStatus is not None:
                self.take_screenshot(page, test_data, f'Verify UC4 Job Execution Status', document)
            else:
                self.take_screenshot(page, test_data, f'FAIL - Verify UC4 Job Execution Status', document)
    def test_load_table_names(self, test_data):
        TestGoogleCloudWorkflow.lstTotalTables = test_data['TableName'].split(";")
        TestGoogleCloudWorkflow.lisCurratedTableName = TestGoogleCloudWorkflow.lstTotalTables[0]
        TestGoogleCloudWorkflow.lisHistoryTableName = [intTableNameIterator for intTableNameIterator in TestGoogleCloudWorkflow.lstTotalTables if "hist" in intTableNameIterator]
        TestGoogleCloudWorkflow.lisViewTableName = [intTableNameIterator for intTableNameIterator in TestGoogleCloudWorkflow.lstTotalTables if "view" in intTableNameIterator]
    
    def test_launchGCP(self, google_cloud_helper, test_data, document):
        helper, browser, page = google_cloud_helper
        page.goto(f"https://console.cloud.google.com/welcome?project={test_data['ProjectId']}")
        page.wait_for_load_state("networkidle")
        self.take_screenshot(page, test_data, f'launchGCP', document)
    def test_VerifyWhetherTablesAndViewAvailableinGCP(self, google_cloud_helper, test_data, document):
        helper, browser, page = google_cloud_helper
        lstTotalTables = test_data['TableName'].split(";")
        intTableCount = len(lstTotalTables)
        for intTableCountIterator in range(intTableCount):
            BigQueryHelper(helper.playwright).navigate_to_big_query_studio(page)
            lisTableDetails = ["TableDetails:"]
            strTableVerificationStatus, objTableDetails = BigQueryHelper(helper.playwright).verify_table_exists(page, lstTotalTables[intTableCountIterator], test_data['ProjectId'])
            if strTableVerificationStatus != "FAIL" and strTableVerificationStatus is not None:
                self.take_screenshot(page, test_data, f'Verify Whether TablesAndView' + lstTotalTables[intTableCountIterator] + 'Available in GCP', document)
            else:
                self.take_screenshot(page, test_data, f'FAIL - Verify Whether TablesAndView' + lstTotalTables[intTableCountIterator] + ' Not Available in GCP', document)
            if "curr" in lstTotalTables[intTableCountIterator]:
                lisTableDetailsInnerText = objTableDetails.inner_text().split("\n")
                for intTableDetailsIterator in range(len(lisTableDetailsInnerText)):
                    if "Primary key(s)" in lisTableDetailsInnerText[intTableDetailsIterator]:
                        lisTableDetails.append("Primary key(s)" + ":" + lisTableDetailsInnerText[intTableDetailsIterator+1])
                    elif "Partitioned by" in lisTableDetailsInnerText[intTableDetailsIterator]:
                        lisTableDetails.append("Partitioned by"+ ":" + lisTableDetailsInnerText[intTableDetailsIterator+1])
                    elif "Partitioned on field" in lisTableDetailsInnerText[intTableDetailsIterator]:
                        lisTableDetails.append("Partitioned on field"+ ":" + lisTableDetailsInnerText[intTableDetailsIterator+1])
                    elif "Partition expiry" in lisTableDetailsInnerText[intTableDetailsIterator]:
                        lisTableDetails.append("Partition expiry" + ":" + lisTableDetailsInnerText[intTableDetailsIterator+1])
                        break
                self.write_output_to_csv(lisTableDetails,TestGoogleCloudWorkflow.output_csv_filename)
    def test_Verify_schema_of_gcp_table(self, google_cloud_helper, test_data, document):
        helper, browser, page = google_cloud_helper
        SchemaValidationStatus = SchemaValidation(helper.playwright).VerifySchemaValidation(page, test_data['Mapping_doc_link'], test_data['ProjectId'], (test_data['TableName'].split(';'))[0])
        print(SchemaValidationStatus)
        if SchemaValidationStatus == "PASS":
            document.add_heading(f'Verify Schema For '+test_data['TableName'], level=2)
            paragraph = document.add_paragraph('SchemaValidation_TestResult_Link : ')
            run = paragraph.add_run(test_data['Mapping_doc_link'])
            run.font.underline = True
            run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        else:
            document.add_heading(f'FAIL - Verify Schema For '+test_data['TableName'], level=2)
            paragraph = document.add_paragraph('SchemaValidation_TestResult_Link : ')
            run = paragraph.add_run(test_data['Mapping_doc_link'])
            run.font.underline = True
            run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        assert(SchemaValidationStatus == "PASS")
    def test_NavigateToInboundBucket(self, google_cloud_helper, test_data, document):
        helper, browser, page = google_cloud_helper
        page.goto(f"https://console.cloud.google.com/welcome?project={test_data['ProjectId']}")
        CloudStorageHelper(helper.playwright).navigate_to_gcp_cloud_storage(page)
        strInboundBucketVerificationStatus = CloudStorageHelper(helper.playwright).GCP_Buckets_navigate_and_click_storage_bucket(page, test_data['InboundBucket'])
        if strInboundBucketVerificationStatus != "FAIL" and strInboundBucketVerificationStatus is not None:
            self.take_screenshot(page, test_data, f'NavigateToInboundBucket-'+test_data['InboundBucket'], document)
        else:
            self.take_screenshot(page, test_data, f'FAIL - NavigateToInboundBucket-'+test_data['InboundBucket'], document)
    
    def test_fetch_inboundfiles_from_bucket(self, google_cloud_helper, test_data, document):
        helper, browser, page = google_cloud_helper
        if "." not in test_data['InboundFileName'] and TestGoogleCloudWorkflow.strUC4JobEndTime is not None:
            strTestData = test_data['InboundFileName'] + "_" + TestGoogleCloudWorkflow.strUC4JobEndTime
        else:
            strTestData = test_data['InboundFileName']
        TestGoogleCloudWorkflow.strInboundFileName = CloudStorageHelper(helper.playwright).GCP_Buckets_Details_fetch_recent_file(page, "Created", strTestData)
        if TestGoogleCloudWorkflow.strInboundFileName != "FAIL" and TestGoogleCloudWorkflow.strInboundFileName is not None:
            self.take_screenshot(page, test_data,f'FetchInboundFiles-' +  TestGoogleCloudWorkflow.strInboundFileName  +  'From' +  test_data['InboundBucket'], document)
        else:
            self.take_screenshot(page, test_data, 'FAIL - FetchInboundFiles-' + test_data['InboundFileName'] + 'From' + test_data['InboundBucket'], document)
    
    def test_dag_run_status(self,google_cloud_helper,test_data,document):
        helper, browser, page = google_cloud_helper
        strActualFileName = TestGoogleCloudWorkflow.strInboundFileName
        if "." not in test_data['InboundFileName'] and strActualFileName is None:
            strActualFileName = test_data['InboundFileName'] + "_" + datetime.datetime.now().strftime("%Y%m%d%H")
        else:
            strActualFileName = test_data['InboundFileName']
        dag_status=DagComposerHelper(helper.playwright).Navigate_to_Dag_url_and_check_the_Dag_Status(page,test_data['DagComposerID'],strActualFileName)
        print(f"this is dag status:{dag_status}")
        if dag_status != "FAIL" and dag_status is not None:
             self.take_screenshot(page, test_data,'Checking the dag status in the environment' + test_data['DagComposerID'], document)
        else:
            self.take_screenshot(page, test_data,'FAIL - file is not processed through the dag',document)
        page.goto(f"https://console.cloud.google.com/welcome?project={test_data['ProjectId']}")
    
    def test_verify_files_in_processed_bucket(self, google_cloud_helper, test_data, document):
        helper, browser, page = google_cloud_helper
        strActualFileName = TestGoogleCloudWorkflow.strInboundFileName
        if "." not in test_data['InboundFileName'] and strActualFileName is None:
            strActualFileName = test_data['InboundFileName'] + "_" + datetime.datetime.now().strftime("%Y%m%d%H")
        else:
            strActualFileName = test_data['InboundFileName']
        CloudStorageHelper(helper.playwright).navigate_to_gcp_cloud_storage(page)
        CloudStorageHelper(helper.playwright).GCP_Buckets_navigate_and_click_storage_bucket(page, test_data['ProcessedBucket'])
        TestGoogleCloudWorkflow.strActualProcessedFileName = CloudStorageHelper(helper.playwright).GCP_Buckets_Details_fetch_recent_file(page, "Created", strActualFileName)
        print(TestGoogleCloudWorkflow.strActualProcessedFileName)
        if TestGoogleCloudWorkflow.strActualProcessedFileName != "FAIL" and TestGoogleCloudWorkflow.strActualProcessedFileName is not None:
            self.take_screenshot(page, test_data, f'VerifyFilesInProcessedBucket-'+strActualFileName, document)
        else:
            self.take_screenshot(page, test_data, f'FAIL - VerifyFilesInProcessedBucket-'+test_data['ProcessedBucket'], document)
    
    def test_verify_configfile_exists(self, google_cloud_helper, test_data, document):
        helper, browser, page = google_cloud_helper
        BigQueryHelper(helper.playwright).navigate_to_big_query_studio(page)
        CloudStorageHelper(helper.playwright).navigate_to_gcp_cloud_storage(page)
        CloudStorageHelper(helper.playwright).GCP_Buckets_navigate_and_click_storage_bucket(page, test_data['ConfigBucket'])
        strConfigFileValidationStatus = CloudStorageHelper(helper.playwright).GCP_Buckets_Details_fetch_recent_file(page, "Created", test_data['ConfigFileName'])
        if strConfigFileValidationStatus != "FAIL" and strConfigFileValidationStatus is not None:
            self.take_screenshot(page, test_data, f'VerifyConfigFileExists-'+test_data['ConfigFileName'], document)
        else:
            self.take_screenshot(page, test_data, f'FAIL - VerifyConfigFileExists-'+test_data['ConfigFileName'], document)
    def test_verify_audit_table_Validations(self, google_cloud_helper, test_data, document):
        helper, browser, page = google_cloud_helper
        strAuditTableSqlQuery = f'Select * from' + "`" + test_data['ProjectId'] +'.adp_control_data.audit_load_table` where source_file =' +"'"+ str(TestGoogleCloudWorkflow.strActualProcessedFileName) +"'"
        BigQueryHelper(helper.playwright).navigate_to_big_query_studio(page)
        strAuditTableValidationStatus = BigQueryHelper(helper.playwright).run_and_verify_query(page,strAuditTableSqlQuery)
        try:
            dfAuditQueryResults = GCP_BigQuery_Query_Library.GoogleCloudBigQueryHelper().bq_get_table_data(strAuditTableSqlQuery)
        except Exception as e:
            print(f"run_and_verify_query '{strAuditTableSqlQuery}': {e}")
        if strAuditTableValidationStatus != "FAIL" and strAuditTableValidationStatus is not None:
            TestGoogleCloudWorkflow.strJobStartTime = dfAuditQueryResults['job_start_dttm'].iloc[0]
            TestGoogleCloudWorkflow.strJobEndTime = dfAuditQueryResults['job_end_dttm'].iloc[0]
            TestGoogleCloudWorkflow.strNumberOfRecords = dfAuditQueryResults['num_of_rec_loaded'].iloc[0]
            self.take_screenshot(page, test_data, f'VerifyAuditTableValidations-'+test_data['InboundFileName'], document)
        else:
            self.take_screenshot(page, test_data, f'FAIL - VerifyAuditTableValidations-'+test_data['InboundFileName'], document)
    def test_VerifyDataLoadedintoCurratedTable(self, google_cloud_helper, test_data, document):
        helper, browser, page = google_cloud_helper
        strCurratedTableCountSqlQuery = f'Select Count(*) from' + "`" + test_data['ProjectId'] + '.' + TestGoogleCloudWorkflow.lisCurratedTableName +"`" + " " + "where processing_dttm" + " " + "between"  + "'"+str(TestGoogleCloudWorkflow.strJobStartTime)+"'" + "and" + "'" +str(TestGoogleCloudWorkflow.strJobEndTime)+"'"
        strCurratedTableCountValidationStatus = BigQueryHelper(helper.playwright).run_and_verify_query(page, strCurratedTableCountSqlQuery)
        dfCurratedTableQueryResults = GCP_BigQuery_Query_Library.GoogleCloudBigQueryHelper().bq_get_table_data(strCurratedTableCountSqlQuery)
        if strCurratedTableCountValidationStatus != "FAIL" and strCurratedTableCountValidationStatus is not None:
            if int(dfCurratedTableQueryResults['f0_'].iloc[0]) == TestGoogleCloudWorkflow.strNumberOfRecords:
                self.take_screenshot(page, test_data, f'VerifyDataLoadedintoCurratedTable-'+test_data['InboundFileName'], document)
            else:
                self.take_screenshot(page, test_data, f'FAIL - VerifyDataLoadedintoCurratedTable-'+test_data['InboundFileName'], document)
        else:
            self.take_screenshot(page, test_data, f'FAIL - VerifyDataLoadedintoCurratedTable-'+test_data['InboundFileName'], document)
    def test_VerifyDataLoadedintoHistoryTable(self, google_cloud_helper, test_data, document):
        helper, browser, page = google_cloud_helper
        strHistoryTableCountSqlQuery = f'Select Count(*) from' + "`" + test_data['ProjectId'] + '.' + TestGoogleCloudWorkflow.lisHistoryTableName[0] +"`" + " " + "where processing_dttm" + " " + "between"  + "'"+str(TestGoogleCloudWorkflow.strJobStartTime)+"'" + "and" + "'" +str(TestGoogleCloudWorkflow.strJobEndTime)+"'"
        strHistoryTableValidationStatus = BigQueryHelper(helper.playwright).run_and_verify_query(page, strHistoryTableCountSqlQuery)
        dfHistoryTableQueryResults = GCP_BigQuery_Query_Library.GoogleCloudBigQueryHelper().bq_get_table_data(strHistoryTableCountSqlQuery)
        if strHistoryTableValidationStatus != "FAIL" and strHistoryTableValidationStatus is not None:
            if int(dfHistoryTableQueryResults['f0_'].iloc[0]) == TestGoogleCloudWorkflow.strNumberOfRecords:
                self.take_screenshot(page, test_data, f'VerifyDataLoadedintoHistoryTable-'+test_data['InboundFileName'], document)
            else:
                self.take_screenshot(page, test_data, f'FAIL - VerifyDataLoadedintoHistoryTable-'+test_data['InboundFileName'], document)
        else:
            self.take_screenshot(page, test_data, f'FAIL - VerifyDataLoadedintoHistoryTable-'+test_data['InboundFileName'], document)
    def test_VerifyPrimaryKeyAndDeltaColumnsDuplicatesAndNulls(self, google_cloud_helper, test_data, document):
        helper, browser, page = google_cloud_helper
        strActualPrimaryKey = test_data['PrimaryKeys'].replace(";",",")
        lisPrimaryKeys = strActualPrimaryKey.split(",")
        for intPrimaryKeyIterator in range(len(lisPrimaryKeys)):
            if intPrimaryKeyIterator == 0:
                strActualPrimaryKeyNullQuery = f'Select Count(*) from' + "`" + test_data['ProjectId'] + '.' + TestGoogleCloudWorkflow.lisCurratedTableName + "`" + " " + "where" + " " +lisPrimaryKeys[intPrimaryKeyIterator] +" " +"is null"
            else:
                strActualPrimaryKeyNullQuery = strActualPrimaryKeyNullQuery + " " + "and" + " " + lisPrimaryKeys[intPrimaryKeyIterator] +" " + "is null"
        strPrimaryKeyDuplicatesSqlQuery = f'Select Count(*)' +"," + strActualPrimaryKey +" " +'from' + "`" + test_data['ProjectId'] + '.' + TestGoogleCloudWorkflow.lisCurratedTableName + "`" + " " + "group by" + " " +strActualPrimaryKey + " " + "having count(*) > 1"
        strPrimaryKeyNullValidationStatus = BigQueryHelper(helper.playwright).run_and_verify_query(page, strActualPrimaryKeyNullQuery)
        if strPrimaryKeyNullValidationStatus != "FAIL" and strPrimaryKeyNullValidationStatus is not None:
            self.take_screenshot(page, test_data, f'VerifyPrimaryKeyAndDeltaColumnsNulls-'+test_data['PrimaryKeys'], document)
        else:
            self.take_screenshot(page, test_data, f'FAIL - VerifyPrimaryKeyAndDeltaColumnsNulls-'+test_data['PrimaryKeys'], document)
        strPrimaryKeyDuplicateValidationStatus = BigQueryHelper(helper.playwright).run_and_verify_query(page,strPrimaryKeyDuplicatesSqlQuery)
        if strPrimaryKeyDuplicateValidationStatus != "FAIL" and strPrimaryKeyDuplicateValidationStatus is not None:
            self.take_screenshot(page, test_data,f'VerifyPrimaryKeyAndDeltaColumnsDuplicatesAndNulls-' + test_data['PrimaryKeys'],document)
        else:
            self.take_screenshot(page, test_data,f'FAIL - VerifyPrimaryKeyAndDeltaColumnsDuplicatesAndNulls-' + test_data['PrimaryKeys'], document)
    def test_Data_Comparision_File_and_CurrTable(self,test_data, document):
        dfFileData = GoogleCloudBigQueryHelper().bq_read_file_data_buckets(test_data['ProcessedBucket'],test_data['ProcessedFileFolderPath'],str(TestGoogleCloudWorkflow.strActualProcessedFileName))
        lisSchemaFileData = GoogleCloudBigQueryHelper().fetch_schema_names_json(test_data['SchemaBucket'],test_data['SchemaFileFolderPath'],test_data['SchemaFileName'])
        TestGoogleCloudWorkflow.df_Mapped_FileData = GCP_BigQuery_Query_Library.GoogleCloudBigQueryHelper().bq_replace_fileheader_with_schema(dfFileData, lisSchemaFileData)
        strCurratedTableSqlQuery = f'Select * from' + "`" + test_data['ProjectId'] + '.' + TestGoogleCloudWorkflow.lisCurratedTableName + "`" + " " + "where processing_dttm" + " " + "between" + "'" + str(TestGoogleCloudWorkflow.strJobStartTime) + "'" + "and" + "'" + str(TestGoogleCloudWorkflow.strJobEndTime) + "'"
        dfCurratedTableQueryResults = GoogleCloudBigQueryHelper().bq_get_table_data(strCurratedTableSqlQuery)
        if TestGoogleCloudWorkflow.df_Mapped_FileData is not None and dfCurratedTableQueryResults is not None:
            lisPrimaryKeys = [test_data['PrimaryKeys']]
            DataFramesComparision().Compare_DataFrames(TestGoogleCloudWorkflow.df_Mapped_FileData,dfCurratedTableQueryResults,lisPrimaryKeys)
            document.add_heading(f'test_Data_Comparision_File_and_CurrTable' + test_data['TableName'] + ':Pass')
        else:
            document.add_heading(f'test_Data_Comparision_File_and_CurrTable' + test_data['TableName'] + ':Fail')
    def test_Data_Comparision_file_and_HistoryTable(self,test_data, document):
        strHistoryTableSqlQuery = f'Select * from' + "`" + test_data['ProjectId'] + '.' + TestGoogleCloudWorkflow.lisHistoryTableName[0] + "`" + " " + "where processing_dttm" + " " + "between" + "'" + str(TestGoogleCloudWorkflow.strJobStartTime) + "'" + "and" + "'" + str(TestGoogleCloudWorkflow.strJobEndTime) + "'"
        dfHistoryTableQueryResults = GoogleCloudBigQueryHelper().bq_get_table_data(strHistoryTableSqlQuery)
        if TestGoogleCloudWorkflow.df_Mapped_FileData is not None and dfHistoryTableQueryResults is not None:
            lisPrimaryKeys = [test_data['PrimaryKeys']]
            DataFramesComparision().Compare_DataFrames(TestGoogleCloudWorkflow.df_Mapped_FileData,dfHistoryTableQueryResults, lisPrimaryKeys)
            document.add_heading(f'test_Data_Comparision_File_and_HistoryTable' + test_data['TableName'] + ':Pass')
        else:
            document.add_heading(f'test_Data_Comparision_File_and_HistoryTable' + test_data['TableName'] + ':Fail')

    